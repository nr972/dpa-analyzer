#!/usr/bin/env python3
"""
Generate two synthetic sample DPA documents for the DPA Analyzer project.

1. sample_dpa_strong.docx  -- Comprehensive, well-drafted DPA (target score 85-95)
2. sample_dpa_weak.pdf     -- Weak/incomplete DPA (target score 40-60)
"""

import os
from datetime import date

# ---------------------------------------------------------------------------
# 1. STRONG DPA  (DOCX via python-docx)
# ---------------------------------------------------------------------------
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return heading

def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def generate_strong_dpa():
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # -- Title page --
    title = doc.add_heading("DATA PROCESSING AGREEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Between Acme Corporation and DataShield Services Inc."
    )
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Effective Date: 1 January 2025")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Agreement Reference: DPA-ACME-DS-2025-001")
    run.font.size = Pt(11)
    run.italic = True

    doc.add_page_break()

    # ---- TABLE OF CONTENTS (simplified) ----
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Definitions and Interpretation",
        "2. Scope and Purpose of Processing",
        "3. Obligations of the Controller",
        "4. Obligations of the Processor",
        "5. Documented Instructions",
        "6. Confidentiality",
        "7. Security of Processing",
        "8. Sub-Processors",
        "9. Data Subject Rights",
        "10. Personal Data Breach Notification",
        "11. Data Protection Impact Assessments",
        "12. International Data Transfers",
        "13. CCPA / CPRA Provisions",
        "14. Audit Rights",
        "15. Deletion and Return of Data",
        "16. Liability and Indemnification",
        "17. Term and Termination",
        "18. General Provisions",
        "Annex A \u2013 Details of Processing",
        "Annex B \u2013 Technical and Organizational Measures",
        "Annex C \u2013 List of Approved Sub-Processors",
        "Annex D \u2013 Standard Contractual Clauses Reference",
    ]
    for item in toc_items:
        add_body(doc, item)

    doc.add_page_break()

    # =====================================================================
    # RECITALS
    # =====================================================================
    add_heading_styled(doc, "Recitals", level=1)
    add_body(doc, (
        "WHEREAS, Acme Corporation, a corporation organized under the laws of the State of Delaware, "
        "United States, with its principal office at 742 Evergreen Terrace, Springfield, IL 62704 "
        '(hereinafter referred to as the "Controller"), engages DataShield Services Inc., a corporation '
        "organized under the laws of the State of California, United States, with its principal office "
        'at 1600 Amphitheatre Parkway, Mountain View, CA 94043 (hereinafter referred to as the "Processor"), '
        "to process personal data on the Controller's behalf;"
    ))
    add_body(doc, (
        "WHEREAS, the parties wish to ensure that the processing of personal data complies with "
        "Regulation (EU) 2016/679 (General Data Protection Regulation, \"GDPR\"), the California "
        "Consumer Privacy Act as amended by the California Privacy Rights Act (\"CCPA/CPRA\"), and "
        "all other applicable data protection laws;"
    ))
    add_body(doc, (
        "NOW, THEREFORE, the parties agree to the following terms and conditions governing the "
        "processing of personal data."
    ))

    # =====================================================================
    # 1. DEFINITIONS
    # =====================================================================
    add_heading_styled(doc, "1. Definitions and Interpretation", level=1)
    definitions = [
        ('"Applicable Data Protection Law"', "means the GDPR, the CCPA/CPRA, the UK GDPR, the Swiss Federal Act on Data Protection (FADP), and any other applicable data protection or privacy legislation."),
        ('"Controller"', "means Acme Corporation, which determines the purposes and means of the processing of Personal Data."),
        ('"Data Subject"', "means an identified or identifiable natural person to whom the Personal Data relates."),
        ('"EEA"', "means the European Economic Area."),
        ('"Personal Data"', "means any information relating to a Data Subject that is processed by the Processor on behalf of the Controller under this Agreement, as further described in Annex A."),
        ('"Personal Data Breach"', "means a breach of security leading to the accidental or unlawful destruction, loss, alteration, unauthorized disclosure of, or access to, Personal Data transmitted, stored, or otherwise processed."),
        ('"Processing"', "means any operation or set of operations performed on Personal Data, whether or not by automated means, including collection, recording, organization, structuring, storage, adaptation, alteration, retrieval, consultation, use, disclosure by transmission, dissemination, or otherwise making available, alignment, combination, restriction, erasure, or destruction."),
        ('"Processor"', "means DataShield Services Inc., which processes Personal Data on behalf of the Controller."),
        ('"Sub-Processor"', "means any third party appointed by the Processor to process Personal Data on behalf of the Controller."),
        ('"Standard Contractual Clauses" or "SCCs"', "means the standard contractual clauses for the transfer of personal data to third countries adopted by the European Commission pursuant to Decision (EU) 2021/914."),
        ('"Technical and Organizational Measures" or "TOMs"', "means the security measures described in Annex B to this Agreement."),
    ]
    for term, defn in definitions:
        p = doc.add_paragraph()
        run_term = p.add_run(f"{term} ")
        run_term.bold = True
        run_term.font.size = Pt(11)
        run_def = p.add_run(defn)
        run_def.font.size = Pt(11)

    # =====================================================================
    # 2. SCOPE AND PURPOSE
    # =====================================================================
    add_heading_styled(doc, "2. Scope and Purpose of Processing", level=1)
    add_body(doc, (
        "2.1 This Agreement applies to the Processing of Personal Data by the Processor on behalf "
        "of the Controller as described in Annex A (Details of Processing)."
    ))
    add_body(doc, (
        "2.2 The Processor shall process Personal Data only for the purposes specified in Annex A "
        "and in accordance with the Controller's documented instructions, unless required to do so "
        "by European Union or Member State law to which the Processor is subject."
    ))
    add_body(doc, (
        "2.3 The categories of data subjects, types of personal data, and the nature and purpose of "
        "the processing are set out in Annex A."
    ))

    # =====================================================================
    # 3. OBLIGATIONS OF THE CONTROLLER
    # =====================================================================
    add_heading_styled(doc, "3. Obligations of the Controller", level=1)
    add_body(doc, (
        "3.1 The Controller shall ensure that it has a lawful basis for the Processing of Personal "
        "Data and that all necessary consents, notices, and authorizations have been obtained."
    ))
    add_body(doc, (
        "3.2 The Controller shall provide documented instructions to the Processor regarding the "
        "Processing of Personal Data, including the nature, scope, and purpose of the Processing."
    ))
    add_body(doc, (
        "3.3 The Controller shall notify the Processor without undue delay of any changes to "
        "Applicable Data Protection Law that may affect the Processor's obligations under this Agreement."
    ))

    # =====================================================================
    # 4. OBLIGATIONS OF THE PROCESSOR
    # =====================================================================
    add_heading_styled(doc, "4. Obligations of the Processor", level=1)
    add_body(doc, (
        "4.1 The Processor shall process Personal Data only on the basis of documented instructions "
        "from the Controller, including with regard to transfers of Personal Data to a third country "
        "or an international organization, unless required to do so by Union or Member State law; in "
        "such a case, the Processor shall inform the Controller of that legal requirement before "
        "Processing, unless that law prohibits such information on important grounds of public interest."
    ))
    add_body(doc, (
        "4.2 The Processor shall immediately inform the Controller if, in its opinion, an instruction "
        "from the Controller infringes Applicable Data Protection Law."
    ))
    add_body(doc, (
        "4.3 The Processor shall maintain a record of all categories of Processing activities carried "
        "out on behalf of the Controller in accordance with Article 30(2) of the GDPR."
    ))

    # =====================================================================
    # 5. DOCUMENTED INSTRUCTIONS
    # =====================================================================
    add_heading_styled(doc, "5. Documented Instructions", level=1)
    add_body(doc, (
        "5.1 The Controller's instructions for the Processing of Personal Data are set out in this "
        "Agreement, including Annex A. Additional instructions may be issued by the Controller in "
        "writing from time to time."
    ))
    add_body(doc, (
        "5.2 The Processor shall not process Personal Data for any purpose other than as instructed "
        "by the Controller. If the Processor believes that a particular instruction is in violation of "
        "Applicable Data Protection Law, it shall promptly notify the Controller and shall be entitled "
        "to suspend performance of the relevant instruction until the Controller confirms or modifies it."
    ))
    add_body(doc, (
        "5.3 All instructions shall be documented and retained for the duration of this Agreement "
        "and for a period of three (3) years thereafter."
    ))

    # =====================================================================
    # 6. CONFIDENTIALITY
    # =====================================================================
    add_heading_styled(doc, "6. Confidentiality", level=1)
    add_body(doc, (
        "6.1 The Processor shall ensure that all persons authorized to process Personal Data have "
        "committed themselves to confidentiality or are under an appropriate statutory obligation of "
        "confidentiality."
    ))
    add_body(doc, (
        "6.2 The Processor shall ensure that access to Personal Data is limited to those personnel "
        "who need access to perform the services under this Agreement and who have received appropriate "
        "training on data protection requirements."
    ))
    add_body(doc, (
        "6.3 The obligations of confidentiality set out in this Section 6 shall survive the termination "
        "or expiry of this Agreement."
    ))

    # =====================================================================
    # 7. SECURITY OF PROCESSING
    # =====================================================================
    add_heading_styled(doc, "7. Security of Processing", level=1)
    add_body(doc, (
        "7.1 Taking into account the state of the art, the costs of implementation, and the nature, "
        "scope, context, and purposes of Processing, as well as the risk of varying likelihood and "
        "severity for the rights and freedoms of natural persons, the Processor shall implement "
        "appropriate technical and organizational measures to ensure a level of security appropriate "
        "to the risk, as described in Annex B."
    ))
    add_body(doc, "7.2 Such measures shall include, as appropriate:")
    security_measures = [
        "The pseudonymization and encryption of Personal Data, including AES-256 encryption at rest and TLS 1.2 or higher in transit;",
        "The ability to ensure the ongoing confidentiality, integrity, availability, and resilience of processing systems and services;",
        "The ability to restore the availability and access to Personal Data in a timely manner in the event of a physical or technical incident, with a Recovery Time Objective (RTO) of 4 hours and a Recovery Point Objective (RPO) of 1 hour;",
        "A process for regularly testing, assessing, and evaluating the effectiveness of technical and organizational measures for ensuring the security of the Processing, including annual penetration testing and quarterly vulnerability assessments;",
        "Multi-factor authentication (MFA) for all administrative access to systems processing Personal Data;",
        "Intrusion detection and prevention systems (IDS/IPS) monitored 24/7 by a dedicated Security Operations Center (SOC);",
        "Role-based access controls (RBAC) with the principle of least privilege applied to all systems.",
    ]
    for measure in security_measures:
        add_bullet(doc, measure)

    add_body(doc, (
        "7.3 The Processor shall maintain ISO 27001 certification (or an equivalent standard) for all "
        "facilities and systems used to process Personal Data throughout the term of this Agreement."
    ))
    add_body(doc, (
        "7.4 The Processor shall conduct an annual risk assessment and provide a summary of findings "
        "and remediation actions to the Controller upon request."
    ))

    # =====================================================================
    # 8. SUB-PROCESSORS
    # =====================================================================
    add_heading_styled(doc, "8. Sub-Processors", level=1)
    add_body(doc, (
        "8.1 The Controller provides general written authorization for the Processor to engage "
        "Sub-Processors, subject to the conditions set out in this Section 8. The list of currently "
        "approved Sub-Processors is set out in Annex C."
    ))
    add_body(doc, (
        "8.2 The Processor shall notify the Controller in writing at least thirty (30) days in advance "
        "of any intended changes concerning the addition or replacement of Sub-Processors, thereby giving "
        "the Controller the opportunity to object to such changes."
    ))
    add_body(doc, (
        "8.3 The Controller may object to the appointment of a new Sub-Processor by providing written "
        "notice to the Processor within fifteen (15) days of receipt of the notification referred to in "
        "Section 8.2. The objection must be based on reasonable grounds related to data protection. If "
        "the Controller objects, the Processor shall use reasonable efforts to make available an "
        "alternative solution that avoids the use of the objected-to Sub-Processor. If no alternative "
        "is available, either party may terminate this Agreement with respect to the affected Processing "
        "activity upon thirty (30) days' written notice."
    ))
    add_body(doc, (
        "8.4 Where the Processor engages a Sub-Processor, the Processor shall impose on the "
        "Sub-Processor, by way of a written contract, the same data protection obligations as set "
        "out in this Agreement, in particular providing sufficient guarantees to implement appropriate "
        "technical and organizational measures."
    ))
    add_body(doc, (
        "8.5 The Processor shall remain fully liable to the Controller for the performance of the "
        "Sub-Processor's obligations. The Processor shall conduct due diligence on each Sub-Processor "
        "prior to engagement and shall monitor ongoing compliance."
    ))

    # =====================================================================
    # 9. DATA SUBJECT RIGHTS
    # =====================================================================
    add_heading_styled(doc, "9. Data Subject Rights", level=1)
    add_body(doc, (
        "9.1 Taking into account the nature of the Processing, the Processor shall assist the "
        "Controller, by appropriate technical and organizational measures, insofar as this is possible, "
        "for the fulfillment of the Controller's obligation to respond to requests for exercising the "
        "Data Subject's rights under Chapter III of the GDPR, including:"
    ))
    rights = [
        "Right of access (Article 15 GDPR);",
        "Right to rectification (Article 16 GDPR);",
        "Right to erasure / right to be forgotten (Article 17 GDPR);",
        "Right to restriction of processing (Article 18 GDPR);",
        "Right to data portability (Article 20 GDPR);",
        "Right to object (Article 21 GDPR);",
        "Rights related to automated decision-making and profiling (Article 22 GDPR).",
    ]
    for right in rights:
        add_bullet(doc, right)

    add_body(doc, (
        "9.2 The Processor shall notify the Controller without undue delay, and in any event within "
        "five (5) business days, upon receiving any request from a Data Subject. The Processor shall "
        "not respond to a Data Subject request directly unless authorized to do so in writing by the "
        "Controller."
    ))
    add_body(doc, (
        "9.3 The Processor shall provide the Controller with the technical capability to fulfill "
        "Data Subject requests, including the ability to search, extract, modify, and delete Personal "
        "Data within the Processor's systems."
    ))

    # =====================================================================
    # 10. PERSONAL DATA BREACH NOTIFICATION
    # =====================================================================
    add_heading_styled(doc, "10. Personal Data Breach Notification", level=1)
    add_body(doc, (
        "10.1 The Processor shall notify the Controller without undue delay, and in any event within "
        "forty-eight (48) hours after becoming aware of a Personal Data Breach, providing the "
        "Controller with sufficient information to enable compliance with the Controller's obligations "
        "under Article 33 of the GDPR."
    ))
    add_body(doc, "10.2 The notification shall include, at a minimum:")
    breach_items = [
        "A description of the nature of the Personal Data Breach, including where possible the categories and approximate number of Data Subjects concerned and the categories and approximate number of Personal Data records concerned;",
        "The name and contact details of the Processor's data protection officer or other point of contact from whom further information can be obtained;",
        "A description of the likely consequences of the Personal Data Breach;",
        "A description of the measures taken or proposed to be taken by the Processor to address the Personal Data Breach, including, where appropriate, measures to mitigate its possible adverse effects.",
    ]
    for item in breach_items:
        add_bullet(doc, item)

    add_body(doc, (
        "10.3 If the Processor is unable to provide all required information simultaneously, the "
        "information may be provided in phases without further undue delay."
    ))
    add_body(doc, (
        "10.4 The Processor shall cooperate with the Controller and take all reasonable steps to "
        "assist in the investigation, mitigation, and remediation of a Personal Data Breach. The "
        "Processor shall preserve all relevant evidence and logs for a minimum of twelve (12) months "
        "following a Breach."
    ))
    add_body(doc, (
        "10.5 The Processor shall maintain an incident response plan, tested at least annually, and "
        "shall designate a dedicated incident response team available 24/7."
    ))

    # =====================================================================
    # 11. DPIAs
    # =====================================================================
    add_heading_styled(doc, "11. Data Protection Impact Assessments", level=1)
    add_body(doc, (
        "11.1 The Processor shall provide reasonable assistance to the Controller in conducting "
        "data protection impact assessments (DPIAs) and prior consultations with supervisory "
        "authorities where required under Articles 35 and 36 of the GDPR, taking into account "
        "the nature of the Processing and the information available to the Processor."
    ))
    add_body(doc, (
        "11.2 The Processor shall provide the Controller with all information reasonably necessary "
        "to conduct a DPIA, including details of Processing operations, data flows, and security "
        "measures."
    ))

    # =====================================================================
    # 12. INTERNATIONAL DATA TRANSFERS
    # =====================================================================
    add_heading_styled(doc, "12. International Data Transfers", level=1)
    add_body(doc, (
        "12.1 The Processor shall not transfer Personal Data to any country outside the EEA, the "
        "United Kingdom, or Switzerland without the prior written consent of the Controller, unless "
        "required by Union or Member State law."
    ))
    add_body(doc, (
        "12.2 Where transfers of Personal Data to third countries are necessary for the performance "
        "of the services, the parties shall ensure that such transfers are subject to appropriate "
        "safeguards, including:"
    ))
    transfer_items = [
        "The Standard Contractual Clauses (SCCs) adopted by the European Commission pursuant to Decision (EU) 2021/914, as set out in Annex D. For the purposes of the SCCs: Module Two (Controller to Processor) shall apply; the Controller is the data exporter and the Processor is the data importer; the details of Annex I (List of Parties, Description of Transfer, Competent Supervisory Authority) and Annex II (Technical and Organizational Measures) to the SCCs are completed in Annex D of this Agreement;",
        "An adequacy decision by the European Commission under Article 45 of the GDPR, where applicable;",
        "Binding Corporate Rules approved by a competent supervisory authority, where applicable;",
        "The UK International Data Transfer Addendum to the EU SCCs, where Personal Data originating from the United Kingdom is transferred.",
    ]
    for item in transfer_items:
        add_bullet(doc, item)

    add_body(doc, (
        "12.3 The Processor shall conduct a Transfer Impact Assessment (TIA) for each third-country "
        "transfer and provide the results to the Controller upon request."
    ))
    add_body(doc, (
        "12.4 The Processor shall promptly notify the Controller if it becomes aware of any legal "
        "requirement or government access request that may affect the protections afforded to "
        "Personal Data under this Section 12."
    ))

    # =====================================================================
    # 13. CCPA / CPRA PROVISIONS
    # =====================================================================
    add_heading_styled(doc, "13. CCPA / CPRA Provisions", level=1)
    add_body(doc, (
        "13.1 To the extent that the CCPA/CPRA applies to the Processing of Personal Data under "
        "this Agreement, the Processor is designated as a \"Service Provider\" as defined in the "
        "CCPA (Cal. Civ. Code \u00a7 1798.140(ag))."
    ))
    add_body(doc, (
        "13.2 The Processor shall not sell or share (as those terms are defined in the CCPA/CPRA) "
        "the Personal Data received from the Controller."
    ))
    add_body(doc, (
        "13.3 The Processor shall not retain, use, or disclose Personal Data for any purpose other "
        "than for the specific business purpose of performing the services specified in this Agreement, "
        "or as otherwise permitted by the CCPA/CPRA."
    ))
    add_body(doc, (
        "13.4 The Processor shall not combine Personal Data received from the Controller with personal "
        "information that the Processor receives from or on behalf of any other person or collects from "
        "its own interactions with consumers, unless expressly permitted by the CCPA/CPRA."
    ))
    add_body(doc, (
        "13.5 The Processor shall assist the Controller in responding to verifiable consumer requests "
        "under the CCPA/CPRA, including requests to know, requests to delete, and requests to correct, "
        "within the timeframes required by law."
    ))
    add_body(doc, (
        "13.6 The Processor shall comply with all applicable requirements under the CCPA/CPRA and "
        "shall notify the Controller if it can no longer meet its obligations under this Section 13. "
        "In such event, the Controller may direct the Processor to take reasonable steps to stop and "
        "remediate unauthorized use of Personal Data."
    ))
    add_body(doc, (
        "13.7 The Processor grants the Controller the right to take reasonable and appropriate steps "
        "to ensure that the Processor uses Personal Data in a manner consistent with the Controller's "
        "obligations under the CCPA/CPRA."
    ))
    add_body(doc, (
        "13.8 The Processor certifies that it understands and will comply with the restrictions and "
        "obligations set forth in this Section 13 and in the CCPA/CPRA."
    ))

    # =====================================================================
    # 14. AUDIT RIGHTS
    # =====================================================================
    add_heading_styled(doc, "14. Audit Rights", level=1)
    add_body(doc, (
        "14.1 The Processor shall make available to the Controller all information necessary to "
        "demonstrate compliance with the obligations laid down in Article 28 of the GDPR and this "
        "Agreement, and shall allow for and contribute to audits, including inspections, conducted "
        "by the Controller or another auditor mandated by the Controller."
    ))
    add_body(doc, (
        "14.2 The Controller (or its designated auditor) may conduct an on-site or remote audit of "
        "the Processor's facilities, systems, and records related to the Processing of Personal Data "
        "up to once per year, or more frequently if required by Applicable Data Protection Law or "
        "following a Personal Data Breach. The Controller shall provide the Processor with at least "
        "thirty (30) days' prior written notice of any audit."
    ))
    add_body(doc, (
        "14.3 The Processor shall provide reasonable cooperation and access during audits, including "
        "making relevant personnel available for interviews and providing access to relevant systems, "
        "facilities, documentation, and records."
    ))
    add_body(doc, (
        "14.4 In lieu of a Controller-initiated audit, the Processor may, at its option, provide the "
        "Controller with a current SOC 2 Type II report, ISO 27001 certification, or equivalent "
        "third-party audit report covering the systems used to process Personal Data. This does not "
        "limit the Controller's right to conduct its own audit where required by a supervisory "
        "authority or following a Personal Data Breach."
    ))

    # =====================================================================
    # 15. DELETION AND RETURN OF DATA
    # =====================================================================
    add_heading_styled(doc, "15. Deletion and Return of Data", level=1)
    add_body(doc, (
        "15.1 Upon termination or expiry of this Agreement, or upon the Controller's written request, "
        "the Processor shall, at the Controller's choice, delete or return all Personal Data to the "
        "Controller within thirty (30) days, and delete existing copies unless Union or Member State "
        "law requires retention of the Personal Data."
    ))
    add_body(doc, (
        "15.2 The Processor shall provide written certification of deletion to the Controller within "
        "ten (10) business days of completing the deletion."
    ))
    add_body(doc, (
        "15.3 Where return of Personal Data is requested, the Processor shall provide the data in a "
        "commonly used, structured, and machine-readable format (such as CSV, JSON, or XML)."
    ))
    add_body(doc, (
        "15.4 Notwithstanding the foregoing, the Processor may retain Personal Data to the extent "
        "and for the duration required by applicable law, provided that the Processor shall ensure "
        "the confidentiality and security of such retained data and shall not process it for any "
        "other purpose."
    ))

    # =====================================================================
    # 16. LIABILITY AND INDEMNIFICATION
    # =====================================================================
    add_heading_styled(doc, "16. Liability and Indemnification", level=1)
    add_body(doc, (
        "16.1 Each party shall be liable for damage caused by processing that infringes Applicable "
        "Data Protection Law in accordance with Article 82 of the GDPR."
    ))
    add_body(doc, (
        "16.2 The Processor shall indemnify the Controller against all claims, actions, liabilities, "
        "losses, damages, and expenses (including reasonable legal fees) arising from the Processor's "
        "breach of this Agreement or Applicable Data Protection Law."
    ))
    add_body(doc, (
        "16.3 The aggregate liability of the Processor under this Agreement shall not exceed the "
        "total fees paid by the Controller to the Processor in the twelve (12) months preceding the "
        "event giving rise to the liability."
    ))

    # =====================================================================
    # 17. TERM AND TERMINATION
    # =====================================================================
    add_heading_styled(doc, "17. Term and Termination", level=1)
    add_body(doc, (
        "17.1 This Agreement shall commence on the Effective Date and shall continue in force for "
        "the duration of the Processing of Personal Data by the Processor on behalf of the Controller."
    ))
    add_body(doc, (
        "17.2 Either party may terminate this Agreement with ninety (90) days' written notice to the "
        "other party."
    ))
    add_body(doc, (
        "17.3 Either party may terminate this Agreement immediately upon written notice if the other "
        "party materially breaches this Agreement and fails to cure such breach within thirty (30) days "
        "of receipt of written notice thereof."
    ))
    add_body(doc, (
        "17.4 Sections 6 (Confidentiality), 10 (Personal Data Breach), 14 (Audit Rights), "
        "15 (Deletion and Return of Data), and 16 (Liability and Indemnification) shall survive "
        "the termination or expiry of this Agreement."
    ))

    # =====================================================================
    # 18. GENERAL PROVISIONS
    # =====================================================================
    add_heading_styled(doc, "18. General Provisions", level=1)
    add_body(doc, (
        "18.1 Governing Law. This Agreement shall be governed by and construed in accordance with "
        "the laws of the State of Delaware, without regard to its conflict of laws provisions, except "
        "that the data protection provisions shall be interpreted in accordance with the GDPR and other "
        "Applicable Data Protection Law."
    ))
    add_body(doc, (
        "18.2 Entire Agreement. This Agreement (including its Annexes) constitutes the entire agreement "
        "between the parties with respect to the subject matter hereof and supersedes all prior or "
        "contemporaneous agreements, negotiations, and understandings."
    ))
    add_body(doc, (
        "18.3 Amendments. This Agreement may only be amended in writing signed by authorized "
        "representatives of both parties."
    ))
    add_body(doc, (
        "18.4 Severability. If any provision of this Agreement is held to be invalid or unenforceable, "
        "the remaining provisions shall remain in full force and effect."
    ))
    add_body(doc, (
        "18.5 Notices. All notices under this Agreement shall be in writing and sent to the addresses "
        "specified in the Recitals or as otherwise notified in writing by the relevant party."
    ))

    # ---- Signature block ----
    doc.add_page_break()
    add_heading_styled(doc, "Execution", level=1)
    add_body(doc, (
        "IN WITNESS WHEREOF, the parties have executed this Data Processing Agreement as of the "
        "Effective Date."
    ))

    # Signature table
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = "For and on behalf of Acme Corporation (Controller)"
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    table.cell(0, 1).text = "For and on behalf of DataShield Services Inc. (Processor)"
    table.cell(0, 1).paragraphs[0].runs[0].bold = True

    table.cell(1, 0).text = "Signature: ____________________________"
    table.cell(1, 1).text = "Signature: ____________________________"

    table.cell(2, 0).text = "Name: John A. Smith"
    table.cell(2, 1).text = "Name: Sarah L. Chen"

    table.cell(3, 0).text = "Title: Chief Privacy Officer"
    table.cell(3, 1).text = "Title: VP of Data Protection & Compliance"

    table.cell(4, 0).text = "Date: 1 January 2025"
    table.cell(4, 1).text = "Date: 1 January 2025"

    table.cell(5, 0).text = "Email: privacy@acme-corp.example.com"
    table.cell(5, 1).text = "Email: compliance@datashield.example.com"

    # =====================================================================
    # ANNEX A - DETAILS OF PROCESSING
    # =====================================================================
    doc.add_page_break()
    add_heading_styled(doc, "Annex A \u2013 Details of Processing", level=1)

    add_heading_styled(doc, "A.1 Subject Matter and Duration", level=2)
    add_body(doc, (
        "The Processor provides cloud-based customer relationship management (CRM) hosting, data "
        "analytics, and customer support platform services to the Controller. Processing shall "
        "continue for the duration of the service agreement between the parties, initially three (3) "
        "years from the Effective Date."
    ))

    add_heading_styled(doc, "A.2 Nature and Purpose of Processing", level=2)
    add_body(doc, "The nature and purpose of the Processing includes:")
    for item in [
        "Hosting and storage of customer records in the CRM platform;",
        "Processing of customer support tickets and communications;",
        "Analytics and reporting on customer engagement metrics;",
        "Automated backup and disaster recovery of customer data;",
        "Email delivery services for transactional and marketing communications.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "A.3 Categories of Data Subjects", level=2)
    for item in [
        "Customers and prospective customers of the Controller;",
        "Employees and contractors of the Controller who use the CRM platform;",
        "Business contacts and partners of the Controller.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "A.4 Types of Personal Data", level=2)
    for item in [
        "Contact information (name, email address, phone number, postal address);",
        "Account identifiers (customer ID, username);",
        "Transaction and order history;",
        "Customer support communications and records;",
        "Usage data and analytics (login timestamps, feature usage, IP addresses);",
        "Marketing preferences and consent records.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "A.5 Special Categories of Data", level=2)
    add_body(doc, (
        "The Processor is not authorized to process special categories of personal data (Article 9 "
        "GDPR) or personal data relating to criminal convictions and offences (Article 10 GDPR) "
        "unless explicitly instructed in writing by the Controller."
    ))

    # =====================================================================
    # ANNEX B - TECHNICAL AND ORGANIZATIONAL MEASURES
    # =====================================================================
    doc.add_page_break()
    add_heading_styled(doc, "Annex B \u2013 Technical and Organizational Measures", level=1)

    toms = {
        "B.1 Encryption": [
            "All Personal Data encrypted at rest using AES-256.",
            "All data in transit encrypted using TLS 1.2 or higher.",
            "Database-level encryption with customer-managed encryption keys (CMEK) available upon request.",
            "Full disk encryption on all endpoint devices used by personnel with access to Personal Data.",
        ],
        "B.2 Access Controls": [
            "Role-based access control (RBAC) with principle of least privilege.",
            "Multi-factor authentication (MFA) required for all administrative access.",
            "Unique user accounts for all personnel; no shared accounts.",
            "Automated provisioning and de-provisioning linked to HR systems.",
            "Privileged access management (PAM) with session recording for administrative actions.",
            "Quarterly access reviews conducted by the security team.",
        ],
        "B.3 Network Security": [
            "Web Application Firewall (WAF) deployed on all public-facing services.",
            "Network segmentation with dedicated VPCs for customer data environments.",
            "Intrusion Detection and Prevention Systems (IDS/IPS) monitored 24/7.",
            "DDoS protection services in place.",
            "VPN required for remote access to production environments.",
        ],
        "B.4 Physical Security": [
            "Data centers certified to SOC 2 Type II and ISO 27001.",
            "Biometric access controls and 24/7 CCTV monitoring at all facilities.",
            "Visitor management and escort policies.",
            "Environmental controls (fire suppression, temperature, humidity monitoring).",
        ],
        "B.5 Logging and Monitoring": [
            "Centralized logging of all access to systems containing Personal Data.",
            "Security Information and Event Management (SIEM) system with automated alerting.",
            "Audit logs retained for a minimum of twelve (12) months.",
            "Real-time alerting for suspicious activities and policy violations.",
        ],
        "B.6 Business Continuity and Disaster Recovery": [
            "Recovery Time Objective (RTO): 4 hours.",
            "Recovery Point Objective (RPO): 1 hour.",
            "Geographically distributed data replication across at least two regions.",
            "Annual disaster recovery testing with documented results.",
            "Automated failover capabilities for critical services.",
        ],
        "B.7 Personnel Security": [
            "Background checks conducted on all personnel with access to Personal Data.",
            "Mandatory data protection and security awareness training upon hire and annually thereafter.",
            "Acceptable use policies and confidentiality agreements signed by all personnel.",
            "Security incident reporting procedures communicated to all staff.",
        ],
        "B.8 Vulnerability Management": [
            "Annual penetration testing by an independent third party.",
            "Quarterly automated vulnerability scans.",
            "Responsible disclosure and bug bounty program.",
            "Patch management policy requiring critical patches within 72 hours.",
        ],
    }
    for heading, items in toms.items():
        add_heading_styled(doc, heading, level=2)
        for item in items:
            add_bullet(doc, item)

    # =====================================================================
    # ANNEX C - SUB-PROCESSORS
    # =====================================================================
    doc.add_page_break()
    add_heading_styled(doc, "Annex C \u2013 List of Approved Sub-Processors", level=1)
    add_body(doc, "The following Sub-Processors are approved as of the Effective Date:")

    sub_table = doc.add_table(rows=4, cols=4)
    sub_table.style = "Table Grid"
    headers = ["Sub-Processor", "Service Provided", "Location", "Safeguard"]
    for i, h in enumerate(headers):
        sub_table.cell(0, i).text = h
        sub_table.cell(0, i).paragraphs[0].runs[0].bold = True

    sub_data = [
        ["NimbusTech Cloud GmbH", "Cloud infrastructure hosting (IaaS)", "Frankfurt, Germany (EU)", "EU-based; no third-country transfer"],
        ["MailWave Inc.", "Transactional email delivery", "Portland, OR, USA", "EU-US Data Privacy Framework; SCCs"],
        ["SecureVault Ltd.", "Encrypted offsite backup storage", "Zurich, Switzerland", "Swiss adequacy decision; SCCs"],
    ]
    for r, row_data in enumerate(sub_data, start=1):
        for c, cell_data in enumerate(row_data):
            sub_table.cell(r, c).text = cell_data

    # =====================================================================
    # ANNEX D - SCCs REFERENCE
    # =====================================================================
    doc.add_page_break()
    add_heading_styled(doc, "Annex D \u2013 Standard Contractual Clauses Reference", level=1)

    add_heading_styled(doc, "D.1 Applicable Module", level=2)
    add_body(doc, (
        "Module Two (Controller to Processor) of the Standard Contractual Clauses adopted pursuant "
        "to European Commission Decision (EU) 2021/914 shall apply to transfers of Personal Data "
        "from the Controller (data exporter) to the Processor or its Sub-Processors (data importers) "
        "in third countries without an adequate level of data protection."
    ))

    add_heading_styled(doc, "D.2 Annex I to the SCCs \u2013 List of Parties", level=2)
    add_body(doc, "Data Exporter:", bold=True)
    for item in [
        "Name: Acme Corporation",
        "Address: 742 Evergreen Terrace, Springfield, IL 62704, USA",
        "Contact: John A. Smith, Chief Privacy Officer, privacy@acme-corp.example.com",
        "Role: Controller",
    ]:
        add_bullet(doc, item)

    add_body(doc, "Data Importer:", bold=True)
    for item in [
        "Name: DataShield Services Inc.",
        "Address: 1600 Amphitheatre Parkway, Mountain View, CA 94043, USA",
        "Contact: Sarah L. Chen, VP of Data Protection & Compliance, compliance@datashield.example.com",
        "Role: Processor",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "D.3 Annex I \u2013 Description of Transfer", level=2)
    for item in [
        "Categories of Data Subjects: As described in Annex A.3.",
        "Categories of Personal Data: As described in Annex A.4.",
        "Frequency of transfer: Continuous, as part of ongoing service delivery.",
        "Nature and purpose of transfer: As described in Annex A.2.",
        "Retention period: For the duration of the Agreement plus thirty (30) days for deletion, or as required by applicable law.",
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "D.4 Annex I \u2013 Competent Supervisory Authority", level=2)
    add_body(doc, (
        "The competent supervisory authority is the Irish Data Protection Commission (DPC), "
        "in accordance with Clause 13 of the SCCs."
    ))

    add_heading_styled(doc, "D.5 Annex II to the SCCs \u2013 Technical and Organizational Measures", level=2)
    add_body(doc, (
        "The technical and organizational measures applicable to transfers under the SCCs are "
        "as described in Annex B of this Agreement."
    ))

    add_heading_styled(doc, "D.6 Optional Clauses and Elections", level=2)
    for item in [
        "Clause 7 (Docking Clause): Included. Additional parties may accede to the SCCs with the prior written consent of both parties.",
        "Clause 9(a) (Sub-Processor authorization): Option 2 (General written authorization) selected. The Processor shall notify the Controller at least thirty (30) days in advance of any intended addition or replacement of Sub-Processors.",
        "Clause 11 (Redress): The optional language is not included.",
        "Clause 17 (Governing law): The SCCs shall be governed by the laws of Ireland.",
        "Clause 18 (Forum): Disputes shall be resolved before the courts of Ireland.",
    ]:
        add_bullet(doc, item)

    # ---- Save ----
    filepath = os.path.join(OUTPUT_DIR, "sample_dpa_strong.docx")
    doc.save(filepath)
    print(f"[OK] Strong DPA saved to: {filepath}")


# ---------------------------------------------------------------------------
# 2. WEAK DPA  (PDF via fpdf2)
# ---------------------------------------------------------------------------
from fpdf import FPDF


class WeakDPAPDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 5, "Confidential - Globex Industries / QuickCloud Ltd.", align="C")
        self.ln(8)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    def section_heading(self, number, text):
        self.set_font("Helvetica", "B", 13)
        self.set_text_color(30, 30, 80)
        self.ln(4)
        self.cell(0, 8, f"{number}. {text}", new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def body_text(self, text):
        self.set_font("Helvetica", "", 10)
        self.set_text_color(0, 0, 0)
        self.multi_cell(0, 5.5, text)
        self.ln(1.5)

    def sub_heading(self, text):
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(50, 50, 50)
        self.cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")
        self.ln(1)

    def bullet(self, text):
        self.set_font("Helvetica", "", 10)
        self.set_text_color(0, 0, 0)
        x = self.get_x()
        self.cell(8, 5.5, chr(0x2022))  # bullet char
        self.multi_cell(0, 5.5, text)
        self.ln(0.5)


def generate_weak_dpa():
    pdf = WeakDPAPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)

    # ---- Title page ----
    pdf.add_page()
    pdf.ln(40)
    pdf.set_font("Helvetica", "B", 24)
    pdf.set_text_color(30, 30, 80)
    pdf.cell(0, 15, "DATA PROCESSING AGREEMENT", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)
    pdf.set_font("Helvetica", "", 14)
    pdf.set_text_color(60, 60, 60)
    pdf.cell(0, 10, "Between", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(30, 30, 80)
    pdf.cell(0, 10, "Globex Industries", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(60, 60, 60)
    pdf.cell(0, 7, "(\"Controller\")", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 14)
    pdf.cell(0, 10, "and", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(30, 30, 80)
    pdf.cell(0, 10, "QuickCloud Ltd.", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(60, 60, 60)
    pdf.cell(0, 7, "(\"Processor\")", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(12)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 7, "Effective Date: March 15, 2025", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 7, "Ref: GI-QC-DPA-2025", align="C", new_x="LMARGIN", new_y="NEXT")

    # ---- Content starts ----
    pdf.add_page()

    # ----- RECITALS -----
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(30, 30, 80)
    pdf.cell(0, 10, "Recitals", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    pdf.body_text(
        "This Data Processing Agreement (\"Agreement\") is entered into between Globex Industries, "
        "a company incorporated in the United Kingdom with offices at 27 Baker Street, London, "
        "W1U 3BH (\"Controller\"), and QuickCloud Ltd., a company incorporated in the United Kingdom "
        "with offices at 5 Silicon Roundabout, London, EC1V 1JN (\"Processor\")."
    )
    pdf.body_text(
        "The Controller engages the Processor to provide cloud hosting and data processing services. "
        "This Agreement sets out the terms under which the Processor will process personal data on "
        "behalf of the Controller."
    )
    # NOTE: No reference to GDPR Article 28 compliance or other frameworks

    # ----- 1. DEFINITIONS -----
    pdf.section_heading("1", "Definitions")
    pdf.body_text(
        "\"Personal Data\" means any information relating to an identified or identifiable individual "
        "that is processed by the Processor on behalf of the Controller."
    )
    pdf.body_text(
        "\"Processing\" means any operation performed on Personal Data."
    )
    pdf.body_text(
        "\"Data Subject\" means the individual to whom the Personal Data relates."
    )
    # NOTE: Minimal definitions -- missing Sub-Processor, Data Breach, SCCs, etc.

    # ----- 2. SCOPE -----
    pdf.section_heading("2", "Scope of Processing")
    pdf.body_text(
        "The Processor shall process Personal Data as necessary to provide the hosting and data "
        "management services described in the main services agreement between the parties."
    )
    pdf.body_text(
        "The types of personal data include customer names, email addresses, and usage data."
    )
    # NOTE: Vague scope -- no detailed annex, no categories of data subjects

    # ----- 3. PROCESSOR OBLIGATIONS -----
    pdf.section_heading("3", "Obligations of the Processor")
    pdf.body_text(
        "3.1 The Processor shall process Personal Data only in accordance with the Controller's "
        "instructions."
    )
    # NOTE: No mention of "documented" instructions, no mention of what to do
    # if instructions conflict with law

    pdf.body_text(
        "3.2 The Processor shall ensure that its employees who handle Personal Data are subject to "
        "appropriate confidentiality obligations."
    )
    # NOTE: Reasonable but not specific about training or statutory obligations

    pdf.body_text(
        "3.3 The Processor shall implement reasonable measures to protect Personal Data against "
        "unauthorized access, loss, or destruction."
    )
    # WEAK: "reasonable measures" is vague -- no specifics, no reference to
    # Article 32 GDPR, no encryption standards, no ISO certification

    # ----- 4. SUB-PROCESSORS -----
    pdf.section_heading("4", "Sub-Processors")
    pdf.body_text(
        "The Processor may engage third-party sub-processors to assist in the provision of "
        "services. The Processor shall ensure that any sub-processor is bound by data protection "
        "obligations no less protective than those set out in this Agreement."
    )
    # WEAK: No notification mechanism, no objection rights, no list of
    # current sub-processors, no advance notice requirement

    # ----- 5. DATA SECURITY -----
    pdf.section_heading("5", "Data Security")
    pdf.body_text(
        "The Processor shall maintain reasonable technical and organizational measures to protect "
        "Personal Data. These measures shall be appropriate to the nature of the data processed."
    )
    pdf.body_text(
        "The Processor shall use commercially reasonable efforts to maintain the security of its "
        "systems."
    )
    # WEAK: Entirely vague -- no specifics on encryption, access controls,
    # penetration testing, certifications, etc.

    # ----- 6. DATA BREACH -----
    pdf.section_heading("6", "Data Breach Notification")
    pdf.body_text(
        "In the event of a data breach affecting Personal Data processed under this Agreement, "
        "the Processor shall promptly notify the Controller."
    )
    pdf.body_text(
        "The notification shall include a general description of the breach and, where possible, "
        "the categories of data affected."
    )
    # WEAK: No specific timeframe (just "promptly"), minimal content
    # requirements, no mention of cooperation, evidence preservation, or
    # incident response planning

    # ----- 7. DATA SUBJECT RIGHTS -----
    pdf.section_heading("7", "Data Subject Rights")
    pdf.body_text(
        "The Processor shall, where reasonably practicable, assist the Controller in responding "
        "to requests from Data Subjects."
    )
    # WEAK: Vague, no specifics on which rights, no timeframes, "where
    # reasonably practicable" is a significant qualifier

    # ----- 8. DATA TRANSFERS -----
    pdf.section_heading("8", "International Data Transfers")
    pdf.body_text(
        "The Processor may store and process Personal Data in any country where the Processor or "
        "its sub-processors maintain facilities."
    )
    pdf.body_text(
        "The Processor shall take steps to ensure that Personal Data is adequately protected "
        "in accordance with applicable law."
    )
    # WEAK: No transfer mechanism specified (no SCCs, no adequacy decisions,
    # no TIA), allows unrestricted international transfers, no Annex I or II

    # ----- 9. TERM AND TERMINATION -----
    pdf.section_heading("9", "Term and Termination")
    pdf.body_text(
        "This Agreement shall remain in effect for the duration of the services agreement between "
        "the parties."
    )
    pdf.body_text(
        "Either party may terminate this Agreement with sixty (60) days' written notice."
    )
    # NOTE: No provisions for data deletion/return on termination -- critical gap

    # ----- 10. RETURN/DELETION OF DATA -----
    pdf.section_heading("10", "Data Retention")
    pdf.body_text(
        "Upon termination of this Agreement, the Processor shall delete or return Personal Data "
        "to the Controller upon request."
    )
    # WEAK: No timeframe for deletion, no certification of deletion,
    # no format specification for return, "upon request" is passive

    # ----- 11. GENERAL -----
    pdf.section_heading("11", "General Provisions")
    pdf.body_text(
        "11.1 This Agreement shall be governed by the laws of England and Wales."
    )
    pdf.body_text(
        "11.2 This Agreement constitutes the entire understanding between the parties regarding "
        "the processing of Personal Data."
    )
    pdf.body_text(
        "11.3 This Agreement may be amended by mutual written consent of both parties."
    )

    # NOTE: No audit rights section at all
    # NOTE: No CCPA/CPRA provisions
    # NOTE: No DPIA assistance provisions
    # NOTE: No liability or indemnification provisions

    # ---- Signature block ----
    pdf.ln(10)
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(30, 30, 80)
    pdf.cell(0, 8, "Signatures", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(90, 6, "For Globex Industries (Controller):")
    pdf.cell(90, 6, "For QuickCloud Ltd. (Processor):", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(3)
    pdf.set_font("Helvetica", "", 10)

    sig_lines = [
        ("Signature: ____________________", "Signature: ____________________"),
        ("Name: David R. Thompson", "Name: Emily J. Park"),
        ("Title: Head of Legal", "Title: Operations Manager"),
        ("Date: 15 March 2025", "Date: 15 March 2025"),
    ]
    for left, right in sig_lines:
        pdf.cell(90, 6, left)
        pdf.cell(90, 6, right, new_x="LMARGIN", new_y="NEXT")

    # ---- Save ----
    filepath = os.path.join(OUTPUT_DIR, "sample_dpa_weak.pdf")
    pdf.output(filepath)
    print(f"[OK] Weak DPA saved to: {filepath}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    generate_strong_dpa()
    generate_weak_dpa()
    print("\nDone. Both sample DPA documents generated successfully.")
