"""Tests for document parsing and clause segmentation."""

from pathlib import Path

import pytest

from app.services.parser import (
    Clause,
    ParsedDocument,
    _is_likely_heading,
    _segment_clauses,
    parse_document,
)


# ---------------------------------------------------------------------------
# Helper to create a synthetic .docx for testing
# ---------------------------------------------------------------------------


def _create_sample_docx(path: Path) -> None:
    """Create a minimal synthetic DPA Word document."""
    import docx

    doc = docx.Document()
    doc.add_heading("DATA PROCESSING AGREEMENT", level=1)
    doc.add_paragraph(
        "This Data Processing Agreement (\"DPA\") is entered into between "
        "Acme Corporation (\"Controller\") and Globex Industries (\"Processor\")."
    )

    doc.add_heading("1. Definitions", level=2)
    doc.add_paragraph(
        "\"Personal Data\" means any information relating to an identified or "
        "identifiable natural person."
    )
    doc.add_paragraph(
        "\"Processing\" means any operation performed on personal data, including "
        "collection, recording, storage, and erasure."
    )

    doc.add_heading("2. Scope and Purpose", level=2)
    doc.add_paragraph(
        "The Processor shall process Personal Data only on documented instructions "
        "from the Controller for the purpose of providing the Services."
    )

    doc.add_heading("3. Sub-Processors", level=2)
    doc.add_paragraph(
        "The Processor shall not engage any sub-processor without the prior written "
        "consent of the Controller."
    )

    doc.add_heading("4. Security Measures", level=2)
    doc.add_paragraph(
        "The Processor shall implement appropriate technical and organizational "
        "measures to ensure a level of security appropriate to the risk."
    )

    doc.add_heading("5. Data Breach Notification", level=2)
    doc.add_paragraph(
        "The Processor shall notify the Controller without undue delay after "
        "becoming aware of a personal data breach."
    )

    doc.save(str(path))


def _create_sample_pdf(path: Path) -> None:
    """Create a minimal synthetic DPA PDF document using fpdf2."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "DATA PROCESSING AGREEMENT", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, "This Data Processing Agreement is between Acme Corp and Globex Industries.")

    sections = [
        ("1. DEFINITIONS", '"Personal Data" means any information relating to an identified natural person.'),
        ("2. SCOPE AND PURPOSE", "The Processor shall process Personal Data only on documented instructions."),
        ("3. SUB-PROCESSORS", "No sub-processor shall be engaged without prior written consent."),
        ("4. SECURITY MEASURES", "Appropriate technical and organizational measures shall be implemented."),
        ("5. DATA BREACH NOTIFICATION", "Notification of breach without undue delay."),
    ]
    for heading, body in sections:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, heading, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 6, body)

    pdf.output(str(path))


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestDocxParsing:
    def test_parse_docx(self, tmp_path):
        docx_path = tmp_path / "test_dpa.docx"
        _create_sample_docx(docx_path)

        result = parse_document(docx_path)

        assert isinstance(result, ParsedDocument)
        assert result.metadata["source_type"] == "docx"
        assert "Data Processing Agreement" in result.full_text
        assert len(result.clauses) >= 5  # Title + 5 sections

    def test_docx_clauses_have_text(self, tmp_path):
        docx_path = tmp_path / "test_dpa.docx"
        _create_sample_docx(docx_path)

        result = parse_document(docx_path)

        for clause in result.clauses:
            assert clause.text.strip() != ""
            assert clause.index >= 0


class TestPdfParsing:
    def test_parse_pdf(self, tmp_path):
        pdf_path = tmp_path / "test_dpa.pdf"
        _create_sample_pdf(pdf_path)

        result = parse_document(pdf_path)

        assert isinstance(result, ParsedDocument)
        assert result.metadata["source_type"] == "pdf"
        assert "Data Processing Agreement" in result.full_text or "PROCESSING" in result.full_text
        assert len(result.clauses) >= 1

    def test_pdf_clauses_have_text(self, tmp_path):
        pdf_path = tmp_path / "test_dpa.pdf"
        _create_sample_pdf(pdf_path)

        result = parse_document(pdf_path)

        for clause in result.clauses:
            assert clause.text.strip() != ""


class TestUnsupportedFormat:
    def test_reject_unsupported(self, tmp_path):
        txt_path = tmp_path / "test.txt"
        txt_path.write_text("hello")

        with pytest.raises(ValueError, match="Unsupported file type"):
            parse_document(txt_path)


class TestHeadingDetection:
    def test_all_caps_heading(self):
        assert _is_likely_heading("DEFINITIONS") is True

    def test_numbered_heading(self):
        assert _is_likely_heading("1. Definitions") is True
        assert _is_likely_heading("1.1 Scope of Processing") is True
        assert _is_likely_heading("(a) something") is True

    def test_article_heading(self):
        assert _is_likely_heading("ARTICLE 5") is True
        assert _is_likely_heading("SECTION 3") is True
        assert _is_likely_heading("ANNEX I") is True
        assert _is_likely_heading("SCHEDULE 2") is True

    def test_body_text_not_heading(self):
        assert _is_likely_heading(
            "The processor shall implement appropriate technical measures."
        ) is False

    def test_empty_not_heading(self):
        assert _is_likely_heading("") is False


class TestClauseSegmentation:
    def test_segments_by_heading(self):
        paragraphs = [
            {"text": "1. Definitions", "style": "", "is_heading": True},
            {"text": "Term means something.", "style": "", "is_heading": False},
            {"text": "2. Scope", "style": "", "is_heading": True},
            {"text": "Processing is limited.", "style": "", "is_heading": False},
        ]
        clauses = _segment_clauses(paragraphs)
        assert len(clauses) == 2
        assert clauses[0].heading == "1. Definitions"
        assert "Term means something" in clauses[0].text
        assert clauses[1].heading == "2. Scope"

    def test_no_headings_single_clause(self):
        paragraphs = [
            {"text": "Some body text.", "style": "", "is_heading": False},
            {"text": "More body text.", "style": "", "is_heading": False},
        ]
        clauses = _segment_clauses(paragraphs)
        assert len(clauses) == 1
        assert "Some body text" in clauses[0].text

    def test_empty_input(self):
        assert _segment_clauses([]) == []
