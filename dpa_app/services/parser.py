"""Document parsing for Word and PDF DPA files."""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)

# Regex for detecting numbered section headings (e.g., "1.", "1.1", "2.3.1", "(a)")
_NUMBERED_HEADING_RE = re.compile(
    r"^(?:"
    r"\d+(?:\.\d+)*\.?\s+"       # 1. / 1.1 / 2.3.1.
    r"|\([a-z]\)\s+"             # (a) / (b)
    r"|\([ivxlcdm]+\)\s+"        # (i) / (ii) / (iv)
    r"|[A-Z][\.\)]\s+"           # A. / A)
    r"|ARTICLE\s+\d+"            # ARTICLE 1
    r"|SECTION\s+\d+"            # SECTION 1
    r"|CLAUSE\s+\d+"             # CLAUSE 1
    r"|SCHEDULE\s+\d+"           # SCHEDULE 1
    r"|ANNEX\s+[A-Z\d]+"        # ANNEX I / ANNEX A
    r"|APPENDIX\s+[A-Z\d]+"     # APPENDIX A
    r"|EXHIBIT\s+[A-Z\d]+"      # EXHIBIT A
    r")"
)


@dataclass
class Clause:
    """A logical clause/section extracted from a DPA document."""

    index: int
    heading: str | None
    text: str
    page_or_section: str | None = None


@dataclass
class ParsedDocument:
    """Result of parsing a DPA document."""

    full_text: str
    clauses: list[Clause]
    metadata: dict = field(default_factory=dict)


def parse_document(file_path: Path) -> ParsedDocument:
    """Parse a DPA document from file path. Routes to appropriate parser."""
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return _parse_docx(file_path)
    elif suffix == ".pdf":
        return _parse_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


# ---------------------------------------------------------------------------
# Word (.docx) parsing
# ---------------------------------------------------------------------------


def _parse_docx(file_path: Path) -> ParsedDocument:
    """Parse a Word document using python-docx."""
    import docx

    doc = docx.Document(str(file_path))

    paragraphs: list[dict] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        is_heading = "heading" in style_name.lower()
        paragraphs.append({
            "text": text,
            "style": style_name,
            "is_heading": is_heading,
        })

    full_text = "\n\n".join(p["text"] for p in paragraphs)
    clauses = _segment_clauses(paragraphs)

    return ParsedDocument(
        full_text=full_text,
        clauses=clauses,
        metadata={
            "source_type": "docx",
            "paragraph_count": len(paragraphs),
            "file_name": file_path.name,
        },
    )


# ---------------------------------------------------------------------------
# PDF parsing
# ---------------------------------------------------------------------------


def _parse_pdf(file_path: Path) -> ParsedDocument:
    """Parse a PDF document using pdfplumber."""
    import pdfplumber

    paragraphs: list[dict] = []

    with pdfplumber.open(str(file_path)) as pdf:
        page_count = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            for line in page_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                is_heading = _is_likely_heading(line)
                paragraphs.append({
                    "text": line,
                    "style": f"page_{page_num}",
                    "is_heading": is_heading,
                })

    full_text = "\n\n".join(p["text"] for p in paragraphs)
    clauses = _segment_clauses(paragraphs)

    return ParsedDocument(
        full_text=full_text,
        clauses=clauses,
        metadata={
            "source_type": "pdf",
            "page_count": page_count,
            "paragraph_count": len(paragraphs),
            "file_name": file_path.name,
        },
    )


def _is_likely_heading(text: str) -> bool:
    """Heuristic: detect if a PDF line is likely a section heading."""
    stripped = text.strip()
    if not stripped:
        return False
    # All caps short text is likely a heading
    if stripped.isupper() and len(stripped) < 100:
        return True
    # Numbered section pattern at start
    if _NUMBERED_HEADING_RE.match(stripped):
        # But only if the line is relatively short (headings, not body paragraphs)
        if len(stripped) < 150:
            return True
    return False


# ---------------------------------------------------------------------------
# Clause segmentation
# ---------------------------------------------------------------------------


def _segment_clauses(paragraphs: list[dict]) -> list[Clause]:
    """Segment paragraphs into logical clauses.

    Strategy:
    1. Use heading styles (from Word) or heuristic heading detection (from PDF)
    2. Fall back to numbered section detection
    3. Group consecutive body paragraphs under the preceding heading
    """
    if not paragraphs:
        return []

    clauses: list[Clause] = []
    current_heading: str | None = None
    current_text_parts: list[str] = []
    clause_index = 0

    def _flush() -> None:
        nonlocal clause_index
        if current_text_parts:
            clauses.append(
                Clause(
                    index=clause_index,
                    heading=current_heading,
                    text="\n".join(current_text_parts),
                )
            )
            clause_index += 1

    for para in paragraphs:
        text = para["text"]
        is_heading = para["is_heading"]

        # Detect headings via style or numbered-heading regex
        if is_heading or _NUMBERED_HEADING_RE.match(text):
            _flush()
            current_heading = text
            current_text_parts = [text]
        else:
            current_text_parts.append(text)

    _flush()

    # If no headings were detected, treat the entire document as one clause
    if not clauses:
        full = "\n".join(p["text"] for p in paragraphs)
        clauses = [Clause(index=0, heading=None, text=full)]

    return clauses
