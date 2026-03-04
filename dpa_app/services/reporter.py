"""Report generation for DPA analysis results."""

from __future__ import annotations

import json
import logging
from datetime import UTC, datetime
from pathlib import Path

from dpa_app.config import settings
from dpa_app.models.analysis import AnalysisFinding, DPAAnalysis, RequirementsMatrix

logger = logging.getLogger(__name__)


def generate_report(
    analysis: DPAAnalysis,
    findings: list[AnalysisFinding],
    matrices: list[RequirementsMatrix],
    report_format: str,
) -> Path:
    """Generate a gap analysis report in the specified format.

    Returns the path to the generated report file.
    """
    report_data = _build_report_data(analysis, findings, matrices)
    report_dir = settings.reports_dir
    report_dir.mkdir(parents=True, exist_ok=True)

    base_name = f"dpa_report_{analysis.id}_{analysis.file_id[:8]}"

    if report_format == "json":
        return _generate_json_report(report_data, report_dir, base_name)
    elif report_format == "html":
        return _generate_html_report(report_data, report_dir, base_name)
    elif report_format == "docx":
        return _generate_docx_report(report_data, report_dir, base_name)
    elif report_format == "pdf":
        return _generate_pdf_report(report_data, report_dir, base_name)
    else:
        raise ValueError(f"Unsupported report format: {report_format}")


def _build_report_data(
    analysis: DPAAnalysis,
    findings: list[AnalysisFinding],
    matrices: list[RequirementsMatrix],
) -> dict:
    """Build the common report data structure used by all formatters."""
    findings_data = []
    for f in findings:
        findings_data.append({
            "framework": f.framework,
            "requirement_id": f.requirement_id,
            "requirement_name": f.requirement_name,
            "finding_type": f.finding_type.value if hasattr(f.finding_type, "value") else str(f.finding_type),
            "severity": f.severity.value if hasattr(f.severity, "value") else str(f.severity),
            "matched_clause_text": f.matched_clause_text,
            "explanation": f.explanation,
            "remediation": f.remediation,
            "confidence": f.confidence,
        })

    # Compute summary statistics
    severity_counts: dict[str, int] = {}
    type_counts: dict[str, int] = {}
    framework_counts: dict[str, int] = {}
    non_compliant = []

    for fd in findings_data:
        sev = fd["severity"]
        severity_counts[sev] = severity_counts.get(sev, 0) + 1
        ft = fd["finding_type"]
        type_counts[ft] = type_counts.get(ft, 0) + 1
        fw = fd["framework"]
        framework_counts[fw] = framework_counts.get(fw, 0) + 1
        if ft != "compliant":
            non_compliant.append(fd)

    return {
        "report_title": "DPA Gap Analysis Report",
        "document_name": analysis.original_filename,
        "analysis_id": analysis.id,
        "generated_at": datetime.now(UTC).isoformat(),
        "overall_score": analysis.overall_score,
        "framework_scores": analysis.framework_scores or {},
        "model_used": analysis.model_used,
        "total_findings": len(findings_data),
        "summary": {
            "by_severity": severity_counts,
            "by_type": type_counts,
            "by_framework": framework_counts,
        },
        "matrices_used": [
            {"name": m.name, "framework": m.framework, "version": m.version}
            for m in matrices
        ],
        "findings": findings_data,
        "remediation_checklist": [
            {
                "requirement": f["requirement_name"],
                "framework": f["framework"],
                "severity": f["severity"],
                "remediation": f["remediation"],
            }
            for f in non_compliant
            if f["remediation"]
        ],
    }


# ---------------------------------------------------------------------------
# JSON
# ---------------------------------------------------------------------------


def _generate_json_report(data: dict, report_dir: Path, base_name: str) -> Path:
    path = report_dir / f"{base_name}.json"
    with open(path, "w") as f:
        json.dump(data, f, indent=2, default=str)
    return path


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

_HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
         max-width: 960px; margin: 0 auto; padding: 2rem; color: #1a1a1a; }}
  h1 {{ color: #1e3a5f; border-bottom: 2px solid #1e3a5f; padding-bottom: 0.5rem; }}
  h2 {{ color: #2c5282; margin-top: 2rem; }}
  .score-dashboard {{ display: flex; gap: 1.5rem; flex-wrap: wrap; margin: 1rem 0; }}
  .score-card {{ background: #f7fafc; border: 1px solid #e2e8f0; border-radius: 8px;
                 padding: 1rem 1.5rem; min-width: 180px; text-align: center; }}
  .score-card .value {{ font-size: 2rem; font-weight: bold; }}
  .score-card .label {{ color: #718096; font-size: 0.875rem; }}
  .good {{ color: #38a169; }}
  .moderate {{ color: #d69e2e; }}
  .poor {{ color: #e53e3e; }}
  table {{ width: 100%; border-collapse: collapse; margin: 1rem 0; }}
  th, td {{ padding: 0.5rem 0.75rem; text-align: left; border-bottom: 1px solid #e2e8f0; }}
  th {{ background: #f7fafc; font-weight: 600; }}
  .severity-critical {{ color: #e53e3e; font-weight: bold; }}
  .severity-high {{ color: #dd6b20; }}
  .severity-medium {{ color: #d69e2e; }}
  .severity-low {{ color: #38a169; }}
  .finding-type {{ display: inline-block; padding: 0.15rem 0.5rem; border-radius: 4px;
                   font-size: 0.75rem; font-weight: 600; }}
  .type-compliant {{ background: #c6f6d5; color: #276749; }}
  .type-missing {{ background: #fed7d7; color: #9b2c2c; }}
  .type-deviation {{ background: #fefcbf; color: #975a16; }}
  .type-partial {{ background: #feebc8; color: #975a16; }}
  .remediation {{ background: #ebf8ff; border-left: 3px solid #3182ce;
                  padding: 0.75rem; margin: 0.5rem 0; }}
  footer {{ margin-top: 3rem; padding-top: 1rem; border-top: 1px solid #e2e8f0;
            color: #a0aec0; font-size: 0.75rem; }}
</style>
</head>
<body>
<h1>{title}</h1>
<p><strong>Document:</strong> {document_name}<br>
<strong>Generated:</strong> {generated_at}<br>
<strong>Model:</strong> {model_used}</p>

<h2>Compliance Scores</h2>
<div class="score-dashboard">
  <div class="score-card">
    <div class="value {overall_class}">{overall_score:.0f}</div>
    <div class="label">Overall Score</div>
  </div>
  {framework_cards}
</div>

<h2>Summary</h2>
<p>Total findings: {total_findings}</p>
<table>
<tr><th>Category</th><th>Count</th></tr>
{summary_rows}
</table>

<h2>Findings</h2>
<table>
<tr><th>Requirement</th><th>Framework</th><th>Severity</th><th>Finding</th><th>Explanation</th></tr>
{findings_rows}
</table>

<h2>Remediation Checklist</h2>
{remediation_items}

<footer>Generated by Privacy &amp; DPA Analyzer</footer>
</body>
</html>
"""


def _score_class(score: float | None) -> str:
    if score is None:
        return ""
    if score >= 80:
        return "good"
    if score >= 50:
        return "moderate"
    return "poor"


def _finding_type_class(ft: str) -> str:
    if ft == "compliant":
        return "type-compliant"
    if ft == "missing_provision":
        return "type-missing"
    if ft == "deviation":
        return "type-deviation"
    return "type-partial"


def _generate_html_report(data: dict, report_dir: Path, base_name: str) -> Path:
    # Framework score cards
    cards = []
    for fw, score in data["framework_scores"].items():
        cls = _score_class(score)
        cards.append(
            f'<div class="score-card"><div class="value {cls}">{score:.0f}</div>'
            f'<div class="label">{fw}</div></div>'
        )

    # Summary rows
    summary_rows = ""
    for sev, count in data["summary"]["by_severity"].items():
        summary_rows += f"<tr><td>{sev}</td><td>{count}</td></tr>\n"

    # Findings rows
    findings_rows = ""
    for f in data["findings"]:
        sev_cls = f"severity-{f['severity']}"
        ft_cls = _finding_type_class(f["finding_type"])
        ft_label = f["finding_type"].replace("_", " ")
        findings_rows += (
            f'<tr><td>{_esc(f["requirement_name"])}</td>'
            f'<td>{_esc(f["framework"])}</td>'
            f'<td class="{sev_cls}">{_esc(f["severity"])}</td>'
            f'<td><span class="finding-type {ft_cls}">{_esc(ft_label)}</span></td>'
            f'<td>{_esc(f["explanation"])}</td></tr>\n'
        )

    # Remediation items
    remediation_items = ""
    for item in data["remediation_checklist"]:
        remediation_items += (
            f'<div class="remediation"><strong>{_esc(item["requirement"])}</strong> '
            f'({_esc(item["framework"])}, {_esc(item["severity"])})<br>'
            f'{_esc(item["remediation"])}</div>\n'
        )
    if not remediation_items:
        remediation_items = "<p>No remediations needed — full compliance.</p>"

    html = _HTML_TEMPLATE.format(
        title=data["report_title"],
        document_name=_esc(data["document_name"]),
        generated_at=data["generated_at"],
        model_used=data["model_used"] or "N/A",
        overall_score=data["overall_score"] or 0,
        overall_class=_score_class(data["overall_score"]),
        framework_cards="\n".join(cards),
        total_findings=data["total_findings"],
        summary_rows=summary_rows,
        findings_rows=findings_rows,
        remediation_items=remediation_items,
    )

    path = report_dir / f"{base_name}.html"
    with open(path, "w") as f:
        f.write(html)
    return path


def _esc(text: str | None) -> str:
    """Escape HTML special characters."""
    if text is None:
        return ""
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


# ---------------------------------------------------------------------------
# Word (.docx)
# ---------------------------------------------------------------------------


def _generate_docx_report(data: dict, report_dir: Path, base_name: str) -> Path:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(data["report_title"], level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Document info
    doc.add_paragraph(
        f"Document: {data['document_name']}\n"
        f"Generated: {data['generated_at']}\n"
        f"Model: {data['model_used'] or 'N/A'}"
    )

    # Overall score
    doc.add_heading("Compliance Scores", level=1)
    score_text = f"Overall Score: {data['overall_score']:.0f}/100"
    p = doc.add_paragraph(score_text)
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    score = data["overall_score"] or 0
    if score >= 80:
        run.font.color.rgb = RGBColor(0x38, 0xA1, 0x69)
    elif score >= 50:
        run.font.color.rgb = RGBColor(0xD6, 0x9E, 0x2E)
    else:
        run.font.color.rgb = RGBColor(0xE5, 0x3E, 0x3E)

    for fw, fw_score in data["framework_scores"].items():
        doc.add_paragraph(f"{fw}: {fw_score:.0f}/100")

    # Findings table
    doc.add_heading("Findings", level=1)
    if data["findings"]:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        headers = ["Requirement", "Framework", "Severity", "Finding", "Explanation"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for f in data["findings"]:
            row = table.add_row()
            row.cells[0].text = f["requirement_name"]
            row.cells[1].text = f["framework"]
            row.cells[2].text = f["severity"]
            row.cells[3].text = f["finding_type"].replace("_", " ")
            row.cells[4].text = f["explanation"][:200]
    else:
        doc.add_paragraph("No findings.")

    # Remediation checklist
    doc.add_heading("Remediation Checklist", level=1)
    if data["remediation_checklist"]:
        for item in data["remediation_checklist"]:
            doc.add_paragraph(
                f"{item['requirement']} ({item['framework']}, {item['severity']})",
                style="List Bullet",
            )
            doc.add_paragraph(item["remediation"])
    else:
        doc.add_paragraph("No remediations needed — full compliance.")

    path = report_dir / f"{base_name}.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# PDF (using fpdf2)
# ---------------------------------------------------------------------------


def _generate_pdf_report(data: dict, report_dir: Path, base_name: str) -> Path:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, data["report_title"], new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(4)

    # Document info
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, f"Document: {data['document_name']}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Generated: {data['generated_at']}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Model: {data['model_used'] or 'N/A'}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    # Overall score
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "Compliance Scores", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 24)
    score = data["overall_score"] or 0
    if score >= 80:
        pdf.set_text_color(56, 161, 105)
    elif score >= 50:
        pdf.set_text_color(214, 158, 46)
    else:
        pdf.set_text_color(229, 62, 62)
    pdf.cell(0, 14, f"{score:.0f}/100", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)

    pdf.set_font("Helvetica", "", 11)
    for fw, fw_score in data["framework_scores"].items():
        pdf.cell(0, 7, f"{fw}: {fw_score:.0f}/100", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    # Findings
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "Findings", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "", 9)
    for f in data["findings"]:
        ft_label = f["finding_type"].replace("_", " ")
        line = f"[{f['severity'].upper()}] {f['requirement_name']} - {ft_label}"
        pdf.cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")
        # Truncate explanation for PDF
        explanation = f["explanation"][:300]
        pdf.set_font("Helvetica", "", 8)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 4, explanation)
        pdf.set_font("Helvetica", "", 9)
        pdf.ln(2)

    # Remediation checklist
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "Remediation Checklist", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "", 10)
    if data["remediation_checklist"]:
        for item in data["remediation_checklist"]:
            pdf.set_font("Helvetica", "B", 10)
            req_text = f"{item['requirement']} ({item['framework']}, {item['severity']})"
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 6, req_text)
            pdf.set_font("Helvetica", "", 9)
            rem_text = item["remediation"] or ""
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 5, rem_text)
            pdf.ln(3)
    else:
        pdf.cell(0, 7, "No remediations needed.", new_x="LMARGIN", new_y="NEXT")

    path = report_dir / f"{base_name}.pdf"
    pdf.output(str(path))
    return path
